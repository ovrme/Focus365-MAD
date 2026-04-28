"""Generate a Focus24 project report as Microsoft Word (.docx)."""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


PRIMARY = RGBColor(0x2E, 0x5A, 0xAC)
ACCENT = RGBColor(0x6E, 0x32, 0x9D)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Focus24")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A Cross-Platform Daily Task Manager built with Flutter & Firebase")
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = GREY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project Report\n").bold = True
    meta.add_run("Version 1.0.0\n")
    meta.add_run("Prepared 2026")
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = PRIMARY


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = DARK


def para(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    return p


def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def build():
    doc = Document()
    # Default body font.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc)

    # ---- 1. Executive summary ------------------------------------------------
    h1(doc, "1. Executive Summary")
    para(
        doc,
        "Focus24 is a cross-platform productivity and task-management application "
        "for Android and iOS. The name reflects its core philosophy: stay focused "
        "every day of the year. The app combines daily task tracking, Pomodoro-style "
        "focus sessions, calendar planning, shared collaboration lists, statistics, "
        "notes, and goal streaks into a single product. It supports English and "
        "Khmer (ខ្មែរ) localization, light/dark theming, biometric lock, offline "
        "caching, push reminders, and import/export of user data.",
    )
    para(
        doc,
        "The application is built with Flutter (Dart) on the client side and uses "
        "Google Firebase as its backend: Firebase Authentication for user accounts, "
        "Cloud Firestore for live data sync, and Firebase Cloud Messaging-compatible "
        "local notifications for reminders. State is managed with the Provider package "
        "following an MVVM architecture, and offline data is cached locally with Hive.",
    )

    # ---- 2. Tech stack -------------------------------------------------------
    h1(doc, "2. Technology Stack — Overview")
    table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["UI / Client", "Flutter 3.x (Dart SDK ^3.10)", "Single codebase for Android & iOS"],
            ["State management", "Provider 6.1.5", "MVVM with ChangeNotifier"],
            ["Routing", "go_router 15.x", "Declarative, deep-linkable navigation"],
            ["Auth", "firebase_auth 5.x", "Email/password account management"],
            ["Database", "cloud_firestore 5.6.x", "Real-time document store"],
            ["Local cache", "Hive 2.2.3 + hive_flutter 1.1", "Per-user offline cache & queued ops"],
            ["Notifications", "flutter_local_notifications 18.x", "Task reminders + Pomodoro alerts"],
            ["Charts", "fl_chart 0.70", "Statistics screen visualizations"],
            ["Calendar", "table_calendar 3.x", "Month/week/agenda view"],
            ["Security", "local_auth 2.x", "Biometric app lock"],
            ["Export", "pdf 3.x + open_filex + file_picker", "PDF/CSV/JSON export & restore"],
            ["Localization", "Custom AppLocalizations (en / km)", "Static dual-language resource map"],
            ["Animations", "confetti 0.7", "Celebrate task completion"],
            ["Typography", "google_fonts 6.2", "Battambang, Nokora, Inter"],
            ["Icons / branding", "flutter_launcher_icons 0.13.1", "Generates Android + iOS launcher icons from PNG"],
        ],
    )

    # ---- 2a. Library inventory ----------------------------------------------
    h1(doc, "2a. Library Inventory (Installed Versions)")
    para(
        doc,
        "Every direct dependency declared in pubspec.yaml is listed below with the "
        "exact version resolved in pubspec.lock and a one-line description of how it "
        "is used in Focus24.",
    )
    h2(doc, "Direct runtime dependencies")
    table(
        doc,
        ["Package", "Version", "How it is used"],
        [
            ["flutter", "SDK", "Core widget toolkit."],
            ["flutter_localizations", "SDK", "Material/Cupertino translations for non-English locales."],
            ["cupertino_icons", "1.0.8", "iOS-style icon set."],
            ["go_router", "15.1.2", "All app routes in lib/routes/router.dart."],
            ["provider", "6.1.5", "MultiProvider in main.dart wires 11 ChangeNotifier classes."],
            ["firebase_core", "3.13.x", "Initialised in main() with DefaultFirebaseOptions.currentPlatform."],
            ["firebase_auth", "5.5.x", "Used by AuthService for sign-up / sign-in / reset / sign-out."],
            ["cloud_firestore", "5.6.12", "All persistent data (tasks, projects, labels, notes, shared lists, activities, comments, user profiles)."],
            ["table_calendar", "3.2.0", "CalendarScreen month/week/agenda view with eventLoader markers."],
            ["fl_chart", "0.70.2", "PieChart, BarChart, LineChart on StatisticsScreen."],
            ["flutter_local_notifications", "18.x", "NotificationService schedules due-date & Pomodoro alerts."],
            ["timezone", "0.10.x", "Required by flutter_local_notifications for scheduled times in user's zone."],
            ["flutter_timezone", "4.x", "Detects the device's IANA timezone string."],
            ["hive", "2.2.3", "Local key-value store for cache + offline queue + Pomodoro state."],
            ["hive_flutter", "1.1.0", "Hive.initFlutter() + path-provider integration."],
            ["intl", "0.20.2", "DateFormat for headers, time pickers, charts."],
            ["uuid", "4.5.1", "Generates document ids for new tasks / sub-tasks / projects / labels / notes / reminders."],
            ["share_plus", "11.x", "Sharing tasks and shared-list invite codes via OS share sheet."],
            ["url_launcher", "6.3.1", "Opens URL attachments and external links."],
            ["connectivity_plus", "6.x", "ConnectivityService listens for online/offline transitions."],
            ["local_auth", "2.x", "Biometric prompt for app lock."],
            ["pdf", "3.11.x", "PDF generation in ExportService.exportToPdf()."],
            ["path_provider", "2.x", "Application documents directory for export & backup files."],
            ["open_filex", "4.6.0", "Opens generated PDFs / CSVs / JSON backups from the snackbar."],
            ["file_picker", "8.x", "Selects a JSON file when restoring a backup."],
            ["confetti", "0.7.0", "Confetti burst over the home scaffold when a task is completed."],
            ["google_fonts", "6.2.1", "Battambang / Nokora for Khmer text, Inter for Latin."],
        ],
    )
    h2(doc, "Dev / build-time dependencies")
    table(
        doc,
        ["Package", "Version", "Role"],
        [
            ["flutter_test", "SDK", "Widget testing harness."],
            ["flutter_lints", "6.0.0", "Lints baseline."],
            ["flutter_launcher_icons", "0.13.1", "Generates AndroidManifest mipmaps + iOS AppIcon set from assets/images/logos/logo.png."],
        ],
    )

    # ---- 2b. Firebase configuration -----------------------------------------
    h1(doc, "2b. Firebase Connection Configuration")
    para(
        doc,
        "Firebase configuration is encapsulated in lib/firebase_options.dart "
        "(auto-generated by the FlutterFire CLI). The same project is shared across "
        "Android, iOS, and Web; each platform uses its own API key and App ID.",
    )
    table(
        doc,
        ["Field", "Value"],
        [
            ["Firebase project ID", "listtodov1"],
            ["Storage bucket", "listtodov1.firebasestorage.app"],
            ["Auth domain (Web)", "listtodov1.firebaseapp.com"],
            ["Messaging sender ID", "935176050024"],
            ["Android App ID", "1:935176050024:android:c20eaab6fa5bb9ee0e72dd"],
            ["iOS App ID", "1:935176050024:ios:0ec47c583999d4fe0e72dd"],
            ["iOS bundle ID", "com.example.todolistapp"],
            ["Web App ID", "1:935176050024:web:106b0642a57560dd0e72dd"],
            ["Web measurement ID", "G-7H9Y1ZRRDY"],
        ],
    )
    para(
        doc,
        "On Android, google-services.json lives under android/app/. On iOS, "
        "GoogleService-Info.plist lives under ios/Runner/. Both files are committed "
        "alongside the source. The Hive initialisation runs in parallel with "
        "Firebase via Future.wait, shaving ~300 ms off cold start.",
    )
    h2(doc, "Firestore collections used")
    table(
        doc,
        ["Collection", "Document fields", "Read by"],
        [
            ["users", "displayName, email, photoUrl, createdAt", "AuthService"],
            ["tasks", "All fields from the Task model (see §6)", "TaskService"],
            ["projects", "id, name, colorValue, userId, order, createdAt", "ProjectService"],
            ["labels", "id, name, colorValue, userId", "LabelService"],
            ["categories", "id, name, iconCode, colorValue, userId", "CategoryService"],
            ["notes", "id, title, body, colorValue, pinned, userId, updatedAt", "NoteService"],
            ["sharedLists", "id, name, ownerId, memberIds[], inviteCode, createdAt", "SharedListService"],
            ["invites", "code, listId, listName, ownerId, createdAt", "SharedListService"],
            ["activities", "id, listId, type, actorId, taskId, payload, at", "ActivityService"],
            ["comments", "id, taskId, authorId, body, createdAt", "CommentService"],
            ["userProfiles", "uid, displayName, photoUrl", "UserProfileService"],
        ],
    )

    # ---- 2c. Platform configuration -----------------------------------------
    h1(doc, "2c. Platform Configuration")
    h2(doc, "Android")
    bullet(doc, "compileSdk / targetSdk: 34 (Android 14).")
    bullet(doc, "Minimum SDK: 21 (Android 5.0) — required for flutter_local_notifications.")
    bullet(doc, "Permissions declared in AndroidManifest.xml: INTERNET, RECEIVE_BOOT_COMPLETED, USE_BIOMETRIC, POST_NOTIFICATIONS, SCHEDULE_EXACT_ALARM, USE_FULL_SCREEN_INTENT.")
    bullet(doc, "Launcher icons under android/app/src/main/res/mipmap-* generated by flutter_launcher_icons from assets/images/logos/logo.png.")
    h2(doc, "iOS")
    bullet(doc, "Bundle identifier: com.example.todolistapp (update before App Store release).")
    bullet(doc, "Capabilities: Push Notifications, Background Modes (remote-notification).")
    bullet(doc, "Info.plist keys: NSFaceIDUsageDescription for local_auth, NSUserNotificationUsageDescription for reminders.")
    bullet(doc, "AppIcon set under ios/Runner/Assets.xcassets/AppIcon.appiconset regenerated by flutter_launcher_icons.")

    # ---- 3. Architecture -----------------------------------------------------
    h1(doc, "3. Architecture")
    para(
        doc,
        "Focus24 follows an MVVM (Model-View-ViewModel) pattern. The view layer is "
        "stateless where possible and consumes ChangeNotifier-based view models via "
        "Provider. View models orchestrate one or more services, each of which talks "
        "to a single backing store (Firestore, Hive, or a platform API).",
    )
    h2(doc, "3.1 Folder layout")
    code(
        doc,
        "lib/\n"
        "├── main.dart                # App entry, service init, provider wiring\n"
        "├── app.dart                 # MaterialApp, theme, locale, router\n"
        "├── firebase_options.dart    # Auto-generated Firebase config\n"
        "├── models/                  # Plain Dart data classes (12 models)\n"
        "├── services/                # Firestore/Hive/platform wrappers (16 svc)\n"
        "├── viewmodels/              # ChangeNotifier classes (10 VMs)\n"
        "├── views/\n"
        "│   ├── screens/             # Full-page widgets (20+ screens)\n"
        "│   └── widgets/             # Reusable components\n"
        "├── routes/router.dart       # go_router definition\n"
        "├── theme/app_theme.dart     # Material 3 colour & typography schemes\n"
        "└── utils/\n"
        "    ├── app_localizations.dart    # en/km dictionary + format() helper\n"
        "    └── user_session_bootstrap.dart\n",
    )
    h2(doc, "3.2 Data flow")
    bullet(doc, "Screens read state by calling context.watch<SomeViewModel>().")
    bullet(doc, "User actions invoke methods on the view model (e.g. taskVM.toggleComplete).")
    bullet(doc, "View models delegate to services; services serialize to Firestore (online) and to a per-user Hive box (offline).")
    bullet(doc, "On reconnect, ConnectivityService fires onReconnected, which flushes queued offline operations.")
    bullet(doc, "Hive caches are scoped per user uid to prevent data leakage between accounts.")

    # ---- 4. Startup flow ----------------------------------------------------
    h1(doc, "4. Application Lifecycle")
    h2(doc, "4.1 Cold start (process launch → first screen)")
    bullet(doc, "main() ensures Flutter binding is initialised.")
    bullet(doc, "Firebase.initializeApp() and Hive.initFlutter() run in parallel.")
    bullet(doc, "Legacy global Hive boxes (cached_tasks, pending_operations, streak_data) are deleted once for backwards-compatible migration.")
    bullet(doc, "SettingsViewModel.init() restores theme, locale, biometric, notification, and default-task preferences from disk.")
    bullet(doc, "NotificationService.init() registers the notification channels and the timezone database.")
    bullet(doc, "ConnectivityService starts listening; onReconnected is wired to TaskViewModel.onReconnected for queue flushing.")
    bullet(doc, "MultiProvider wraps the app and exposes 11 view models / services as global state.")
    bullet(doc, "MaterialApp.router boots with the splash route /. SplashScreen plays an entrance animation while authVM.checkAuthState() validates any cached Firebase session.")
    bullet(doc, "After ~1.8 s the router redirects:  no session → /login (or /onboarding if first run); valid session → bootstrapUserSession() loads tasks/projects/labels/notes/categories, then → /home or biometric lock prompt.")
    h2(doc, "4.2 Steady-state runtime")
    bullet(doc, "MainScreen owns a six-tab bottom navigator: Dashboard, Tasks (Home), Projects, Calendar, Statistics, Profile.")
    bullet(doc, "All Firestore reads use snapshots(); UI rebuilds reactively via ChangeNotifier.")
    bullet(doc, "Drafts and offline edits go through TaskViewModel which writes through a Hive cache before pushing to Firestore.")
    h2(doc, "4.3 Background / pause")
    bullet(doc, "PomodoroScreen persists timer state to a Hive box on every state change so timers survive a backgrounded app.")
    bullet(doc, "didChangeAppLifecycleState restores the timer using the recorded timerStartedAt to compute elapsed seconds; if elapsed exceeded the remaining duration the timer auto-completes on resume.")
    bullet(doc, "Notifications continue to fire from flutter_local_notifications even when Focus24 is killed.")
    h2(doc, "4.4 Sign-out / shutdown")
    bullet(doc, "Sign-out calls AuthViewModel.signOut() which forwards to FirebaseAuth.instance.signOut() and clears every view model so subsequent logins start clean.")
    bullet(doc, "Hive boxes are kept on disk; the next sign-in opens the per-uid box and reuses any cached data.")

    # ---- 5. Features --------------------------------------------------------
    h1(doc, "5. Feature Catalogue")
    table(
        doc,
        ["#", "Feature", "What it does"],
        [
            ["1", "Onboarding", "Four-card pager introducing tasks, calendar, statistics, security."],
            ["2", "Authentication", "Email/password sign-up, sign-in, password reset via Firebase Auth."],
            ["3", "Tasks", "Create / edit / complete / delete with priority, category, due date, color, emoji, sub-tasks, attachments, recurrence."],
            ["4", "Recurrence", "Daily / weekly (with weekday picker) / monthly with interval N."],
            ["5", "Reminders", "Per-task multiple reminders; default offset can be set globally."],
            ["6", "Projects", "Group tasks; per-project color, name, kanban view."],
            ["7", "Labels", "Multi-select coloured tags for tasks."],
            ["8", "Categories", "Built-in + custom categories with icon and colour."],
            ["9", "Dashboard", "Greeting, daily quote, streak, productivity score, quick-add templates, today / overdue / upcoming."],
            ["10", "Calendar", "Month grid + agenda or 24-hour timeline view; tap empty hour to create."],
            ["11", "Kanban board", "Drag tasks across To Do / In Progress / Done columns."],
            ["12", "Pomodoro timer", "25-min focus / 5-min break / 15-min long break cycles; per-task session counts."],
            ["13", "Statistics", "Productivity score, current/best streak, completion pie, 30-day trend, day-of-week bars, category & priority breakdown."],
            ["14", "Shared lists", "Multi-user lists with invite codes, member assignment, activity feed."],
            ["15", "Notes", "Coloured sticky notes with pin and grid view."],
            ["16", "Search", "Live cross-collection search across tasks and notes with highlighting."],
            ["17", "Quick add", "Bottom-sheet for one-tap task capture with Today/Tomorrow pills."],
            ["18", "Comments", "Threaded comments on shared tasks."],
            ["19", "Streaks & score", "Streak service counts consecutive active days, computes productivity score 0-100."],
            ["20", "Notifications", "Local notifications for reminders and Pomodoro completion."],
            ["21", "Biometric lock", "Optional fingerprint / face unlock on launch."],
            ["22", "Theme", "Light / dark / system, Material 3 dynamic colour."],
            ["23", "Language", "Full English & Khmer translations across all screens."],
            ["24", "Export", "PDF report, CSV spreadsheet, JSON full backup."],
            ["25", "Restore", "Append JSON backup from file picker or in-app history."],
            ["26", "Offline mode", "Hive cache + queued mutations sync on reconnect."],
            ["27", "Help & FAQ", "Built-in support screen with FAQ and feedback form."],
        ],
    )

    # ---- 6. Models / Services ------------------------------------------------
    h1(doc, "6. Domain Models — Full Field Reference")
    para(
        doc,
        "All models are plain Dart classes with toJson() / fromJson() and "
        "(where the doc lives in Firestore) a fromFirestore(DocumentSnapshot) "
        "factory. Optional fields are only emitted in JSON when set, so legacy "
        "documents round-trip unchanged.",
    )

    h2(doc, "6.1 User")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["uid", "String", "Firebase Auth uid."],
            ["email", "String", "Account email."],
            ["displayName", "String", "Editable in Profile."],
            ["photoUrl", "String?", "Optional avatar URL."],
            ["createdAt", "DateTime", "Account creation time."],
        ],
    )

    h2(doc, "6.2 Task (core entity)")
    table(
        doc,
        ["Field", "Type", "Default", "Notes"],
        [
            ["id", "String", "—", "Server-side or UUID-generated."],
            ["title", "String", "required", "Task title."],
            ["description", "String", "''", "Long description."],
            ["dueDate", "DateTime?", "null", "Optional due date + optional time."],
            ["category", "String", "'General'", "References a TaskCategory.name."],
            ["priority", "int", "2", "1=Low, 2=Medium, 3=High."],
            ["isCompleted", "bool", "false", "Toggled by check button."],
            ["completedAt", "DateTime?", "null", "Set when marked done; used for streak service."],
            ["subTasks", "List<SubTask>", "[]", "Embedded checklist."],
            ["isRecurring", "bool", "false", "Derived flag (true if recurrenceRule != null)."],
            ["recurrenceRule", "RecurrenceRule?", "null", "Daily/weekly/monthly + interval + days-of-week."],
            ["reminders", "List<TaskReminder>", "[]", "Multiple reminders per task."],
            ["userId", "String?", "null", "Owner uid."],
            ["colorValue", "int?", "null", "ARGB int → optional task color stripe."],
            ["emoji", "String?", "null", "Optional unicode glyph prepended in lists."],
            ["projectId", "String?", "null", "FK to projects collection."],
            ["labelIds", "List<String>", "[]", "Many-to-many with labels."],
            ["sharedListId", "String?", "null", "FK to sharedLists when collaborative."],
            ["assigneeId", "String?", "null", "uid of assigned member."],
            ["attachments", "List<String>", "[]", "Plain text or http(s) URL."],
            ["estimatedMinutes", "int?", "null", "Reserved for future analytics."],
            ["completionNote", "String?", "null", "Reserved for future analytics."],
            ["createdAt", "DateTime", "now()", "Insertion timestamp."],
        ],
    )
    para(doc, "Convenience: Task.taskColor returns Color(colorValue!) or null.")

    h2(doc, "6.3 SubTask (embedded in Task)")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["title", "String", "Subtask label."],
            ["isCompleted", "bool", "Checkbox state."],
        ],
    )

    h2(doc, "6.4 RecurrenceRule")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["type", "RecurrenceType", "daily / weekly / monthly / custom."],
            ["interval", "int", "Every N units (default 1)."],
            ["daysOfWeek", "List<int>", "1=Mon..7=Sun, only for weekly."],
        ],
    )

    h2(doc, "6.5 TaskReminder")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["fireAt", "DateTime", "Absolute time to fire the notification."],
            ["offsetMinutesBeforeDue", "int?", "If set, reminder follows due-date changes (e.g. 30 min before)."],
        ],
    )

    h2(doc, "6.6 Project")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["name", "String", "Project label."],
            ["colorValue", "int", "ARGB int; preset palette used in UI."],
            ["userId", "String?", "Owner uid."],
            ["createdAt", "DateTime", "Sort key."],
        ],
    )

    h2(doc, "6.7 Label")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["name", "String", "Label text."],
            ["colorValue", "int", "ARGB int."],
            ["userId", "String?", "Owner uid."],
        ],
    )

    h2(doc, "6.8 TaskCategory")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "Either built-in literal or user UUID."],
            ["name", "String", "User-visible name."],
            ["icon", "IconData", "Material icon."],
            ["color", "Color", "Icon tint."],
            ["userId", "String?", "null for built-ins; uid for custom."],
        ],
    )

    h2(doc, "6.9 Note")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["title", "String", "Optional; fallback to '(Untitled note)'."],
            ["body", "String", "Free text."],
            ["colorValue", "int", "Note background colour."],
            ["pinned", "bool", "Pinned notes float to the top."],
            ["userId", "String", "Owner uid."],
            ["createdAt / updatedAt", "DateTime", "Sort keys."],
        ],
    )

    h2(doc, "6.10 SharedList")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "Firestore doc id."],
            ["name", "String", "List name."],
            ["ownerId", "String", "Creator uid."],
            ["memberIds", "List<String>", "Includes owner."],
            ["inviteCode", "String", "6-char alphanumeric join code."],
            ["createdAt", "DateTime", "Sort key."],
        ],
    )

    h2(doc, "6.11 Invite")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["code", "String", "Doc id (the invite code)."],
            ["listId", "String", "Target list."],
            ["listName", "String", "Cached for the join screen."],
            ["ownerId", "String", "uid of inviter."],
            ["createdAt", "DateTime", ""],
        ],
    )

    h2(doc, "6.12 Activity")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["listId", "String", "Owning shared list."],
            ["type", "String", "task_added / task_completed / member_joined / etc."],
            ["actorId", "String", "uid of the user who triggered it."],
            ["taskId / payload", "Optional", "Context."],
            ["at", "DateTime", "Sort key."],
        ],
    )

    h2(doc, "6.13 Comment")
    table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "String", "UUID."],
            ["taskId", "String", "Parent task."],
            ["authorId", "String", "uid."],
            ["body", "String", "Comment text."],
            ["createdAt", "DateTime", ""],
        ],
    )

    # ---- 6a. Widgets --------------------------------------------------------
    h1(doc, "6a. Widget Inventory")
    h2(doc, "6a.1 Custom widgets (lib/views/widgets/)")
    table(
        doc,
        ["Widget", "Used by", "Purpose"],
        [
            ["AppBackground", "MainScreen (Stack)", "Gradient + soft blobs background under the scaffold."],
            ["QuickAddSheet", "MainScreen FAB", "Bottom sheet for fast task capture with date pills + priority + 'More options'."],
            ["TaskCard", "HomeScreen, DashboardScreen, CalendarScreen, ProjectDetailScreen", "Unified card with checkbox, title, meta row, swipe-to-delete, long-press batch."],
            ["TaskCommentsSection", "AddTaskScreen (shared task)", "Threaded comments stream + composer."],
        ],
    )

    h2(doc, "6a.2 Internal helper widgets (file-private '_Foo')")
    table(
        doc,
        ["Widget", "Location", "Purpose"],
        [
            ["_AppDrawer / _DrawerSection / _DrawerTile / _CountChip", "main_screen.dart", "Side drawer composition."],
            ["_TodayHeader / _SectionHeader / _SectionHeaderDelegate / _TaskDetailSheet", "home_screen.dart", "Header card, pinned section headers, slide-up detail."],
            ["_QuickStatCard / _TemplateChip", "dashboard_screen.dart", "Top stats row + horizontal template chips."],
            ["_StatCard / _LegendItem / _PriorityCard", "statistics_screen.dart", "Stat tiles + pie legend + priority cards."],
            ["_FocusedTaskBanner", "pomodoro_screen.dart", "Banner showing which task the user is focused on."],
            ["_DayTimeline / _AllDayStrip / _HourRow / _TimelineTaskTile", "calendar_screen.dart", "24-hour vertical timeline view."],
            ["_OffsetPick / _RestoreSelection", "settings_screen.dart", "Sentinel types for picker dialogs."],
            ["_EmojiButton / _EmojiPickerSheet / _RemindersRow", "add_task_screen.dart", "Emoji picker + reminders link button."],
            ["_DatePill", "quick_add_sheet.dart", "Reusable selectable chip."],
            ["_EmptyState / _NoteCard / _NoteEditorScreen", "notes_screen.dart", "Notes grid + editor."],
            ["_EmptyHint / _NoResults / _SectionLabel / _TaskResultTile / _NoteResultTile / _Highlighted", "search_screen.dart", "Search UX + highlighted match rendering."],
        ],
    )

    h2(doc, "6a.3 Framework widgets in heavy use")
    table(
        doc,
        ["Category", "Widgets"],
        [
            ["App shell", "MaterialApp.router · Scaffold · AppBar · Drawer · NavigationBar · NavigationDestination · FloatingActionButton"],
            ["Layout", "Column · Row · Stack · Expanded · Padding · SizedBox · Container · Wrap · ListView / ListView.builder · GridView · SingleChildScrollView · CustomScrollView · SliverPersistentHeader · SafeArea · IntrinsicHeight"],
            ["Input", "TextField · TextFormField · Form · GestureDetector · InkWell · Checkbox · Switch · SwitchListTile · IconButton · OutlinedButton · FilledButton · TextButton · SegmentedButton · FilterChip · ChoiceChip · ActionChip · DropdownButtonFormField"],
            ["Feedback / dialog", "Dialog · AlertDialog · SimpleDialog · SnackBar · SnackBarAction · ProgressIndicator · CircularProgressIndicator · LinearProgressIndicator · Tooltip · RefreshIndicator"],
            ["Sheets", "showModalBottomSheet · DraggableScrollableSheet"],
            ["Animation", "AnimationController · Tween · CurvedAnimation · ScaleTransition · FadeTransition · SlideTransition · AnimatedSwitcher · AnimatedRotation"],
            ["Media / shapes", "Image.asset · CircleAvatar · ClipOval · ClipRRect · Icon · Material · DecoratedBox · BoxDecoration"],
            ["Theming", "Theme.of(context) · ThemeData · ColorScheme (Material 3) · TextTheme"],
            ["Provider integration", "Consumer · Selector · context.watch · context.read"],
            ["Calendar / charts", "TableCalendar · PieChart · BarChart · LineChart"],
            ["Misc", "Hero · FutureBuilder · StreamBuilder · ValueListenableBuilder · Builder · StatefulBuilder · WidgetsBinding · WidgetsBindingObserver"],
        ],
    )

    h1(doc, "7. Services Layer")
    table(
        doc,
        ["Service", "Responsibility"],
        [
            ["AuthService", "Firebase Auth wrapper (sign-in, sign-up, reset, sign-out)."],
            ["TaskService", "CRUD + snapshot streams for the tasks collection."],
            ["ProjectService", "Projects CRUD and streams."],
            ["LabelService", "Labels CRUD and streams."],
            ["CategoryService", "User-owned category CRUD; merges with built-ins."],
            ["NoteService", "Notes CRUD."],
            ["SharedListService", "Shared list creation, invite codes, member management."],
            ["CommentService", "Comments on shared tasks."],
            ["ActivityService", "Audit log writes for shared-list mutations."],
            ["UserProfileService", "Cached display-name lookup for shared members."],
            ["CacheService", "Per-user Hive caches & pending-operation queue."],
            ["ConnectivityService", "Streams connectivity; fires onReconnected."],
            ["NotificationService", "Local-notification scheduler for reminders / Pomodoro."],
            ["BiometricService", "local_auth wrapper for biometric prompt."],
            ["StreakService", "Computes daily streaks and 0-100 productivity score."],
            ["ExportService", "PDF / CSV / JSON export, lists in-app backups, restore."],
        ],
    )

    # ---- 8. View models ------------------------------------------------------
    h1(doc, "8. View Models")
    table(
        doc,
        ["ViewModel", "What it exposes"],
        [
            ["AuthViewModel", "User session, sign-in/up/out, password reset, refresh listenable for router."],
            ["TaskViewModel", "Tasks list, active/completed, filters, snapshot subscription, offline queue."],
            ["ProjectViewModel", "Projects collection, active project."],
            ["LabelViewModel", "Labels CRUD passthrough."],
            ["CategoryViewModel", "Merged built-in + user categories."],
            ["NoteViewModel", "Notes CRUD passthrough + pin toggle."],
            ["SettingsViewModel", "Theme, locale, biometric, notifications, default priority/category, default reminder offset, onboarding flag."],
            ["SharedListViewModel", "Owned/member lists, invite acceptance."],
            ["UserProfileViewModel", "Display-name resolver for shared-list members."],
            ["HomeFilterIntent", "Cross-screen filter bus (drawer category/project tap → Home tab)."],
        ],
    )

    # ---- 9. Localization -----------------------------------------------------
    h1(doc, "9. Localization")
    para(
        doc,
        "lib/utils/app_localizations.dart contains a single static map keyed by ISO "
        "language code ('en', 'km') and string identifier. AppLocalizations.of(context).get('key') "
        "returns the localised string with English fallback. A format(key, {params}) helper "
        "performs placeholder substitution (e.g. 'Category · {name}').",
    )
    bullet(doc, "Approximately 230 keys covering every visible UI string.")
    bullet(doc, "Khmer translations use the Battambang/Nokora typefaces via google_fonts.")
    bullet(doc, "Switching locale at runtime is instant; settingsVM.setLocale() persists to Hive.")

    # ---- 10. Issues fixed ----------------------------------------------------
    h1(doc, "10. Issues Encountered & Resolutions")
    para(
        doc,
        "The table below logs every defect discovered during the latest review pass "
        "and the corrective action taken.",
    )
    table(
        doc,
        ["#", "Issue", "Root cause", "Fix"],
        [
            [
                "1",
                "~80 hardcoded English strings remained when Khmer mode was active (drawer, chips, section headers, dialogs, etc.).",
                "Screens were authored before AppLocalizations was extended; raw 'Text(\"...\")' calls leaked English literals.",
                "Added ~80 new key pairs to AppLocalizations (en & km) plus a format() helper, then replaced every literal across 10 screens (main, home, pomodoro, settings, add_task, quick_add, notes, search, calendar, statistics).",
            ],
            [
                "2",
                "Dashboard projects strip showed 'BOTTOM OVERFLOWED BY 4.0 PIXELS' yellow-and-black banner.",
                "Fixed SizedBox(height: 80) was 4 px shorter than the inner Column (10 px circle + 8 px gap + 14 sp text + 12 px padding ×2), especially with Khmer tall diacritics.",
                "Raised height to 92, reduced inner gap to 6, added mainAxisSize.min and maxLines:1 to both texts.",
            ],
            [
                "3",
                "Launcher icon on device still showed the default Flutter logo instead of Focus24 logo.",
                "flutter_launcher_icons was configured but never installed/run.",
                "Added the package, ran flutter pub get, regenerated icons (legacy + adaptive), then iterated config to remove adaptive shrink so the logo uses full width/height. Documented the required flutter clean / re-install step.",
            ],
            [
                "4",
                "Splash screen showed the logo at ~20 % of its actual size inside a coloured circle.",
                "Earlier design wrapped the logo in a 128×128 Container with primaryContainer background and circular ClipOval.",
                "Removed background container and ClipOval; logo now renders 260×260 with BoxFit.contain and FilterQuality.high.",
            ],
            [
                "5",
                "Tapping an attachment row in Add Task did nothing.",
                "The ListTile rendering attachments had no onTap handler.",
                "Added _openAttachment(value): launches URLs via url_launcher; copies plain-text attachments to clipboard with snackbar feedback. URL rows are now styled with primary colour + underline so they look tappable.",
            ],
            [
                "6",
                "User wanted to export all account emails / names / passwords.",
                "Misconception: Firebase passwords are unrecoverable salted hashes; mobile SDK cannot list users.",
                "Created scripts/export-users (Node + Firebase Admin SDK) that exports uid, email, displayName, providers, timestamps to UTF-8 BOM CSV. Clearly states passwords cannot be exported. Friendly error if service-account.json missing.",
            ],
            [
                "7",
                "App name TaskMaster Pro renamed to Focus24.",
                "Brand rename.",
                "Updated appName key in localisation, MaterialApp title, About dialog, biometric prompt, PDF report header, JSON backup metadata, shared-list invite share text, splash subtitle.",
            ],
            [
                "8",
                "Cross-account data leakage risk in Hive caches.",
                "Original boxes were globally named (cached_tasks, pending_operations, streak_data).",
                "Existing main.dart already performs a one-time delete migration and opens per-uid boxes for new users — documented as part of cold-start flow.",
            ],
            [
                "9",
                "Pomodoro timer reset whenever the user backgrounded the app.",
                "Initial implementation did not persist state.",
                "PomodoroScreen now uses Hive (pomodoro_state box) to save totalSeconds, remainingSeconds, isRunning, isBreak, completedPomodoros, timerStartedAt, selectedTaskId, taskCounts; didChangeAppLifecycleState restores and adjusts for elapsed time.",
            ],
        ],
    )

    # ---- 11. Security & privacy ----------------------------------------------
    h1(doc, "11. Security & Privacy")
    bullet(doc, "Authentication handled exclusively by Firebase; passwords never touch the application code.")
    bullet(doc, "Firestore rules restrict each document to its owner uid; shared-list documents check memberIds membership.")
    bullet(doc, "Optional biometric lock (Android fingerprint, iOS Face ID / Touch ID) gates the home screen.")
    bullet(doc, "Local Hive caches are scoped per user uid; sign-out leaves the box on disk but it is only re-opened when the same user signs back in.")
    bullet(doc, "Export functionality is opt-in; users own their data and can take it as PDF/CSV/JSON.")
    bullet(doc, "scripts/export-users/.gitignore excludes service-account.json and users.csv so admin credentials never enter version control.")

    # ---- 12. Testing & quality -----------------------------------------------
    h1(doc, "12. Testing & Quality Assurance")
    bullet(doc, "Static analysis: flutter analyze ran clean after all localisation refactors.")
    bullet(doc, "Linting: flutter_lints 6.x baseline ruleset.")
    bullet(doc, "Manual smoke tests: language toggle, theme toggle, sign-in / sign-up / reset, task CRUD, recurrence, reminders, projects, kanban DnD, calendar timeline, Pomodoro background survival, export PDF/CSV/JSON, restore JSON, biometric lock.")

    # ---- 13. Future work -----------------------------------------------------
    h1(doc, "13. Recommended Next Steps")
    bullet(doc, "Replace literal greeting / motivational quotes with localisation keys so Khmer users see translated quotes.")
    bullet(doc, "Migrate from JSON-backup ad-hoc restore to Firestore Cloud Functions for safer schema upgrades.")
    bullet(doc, "Add widget tests for the core flows (auth gate, task CRUD, recurrence).")
    bullet(doc, "Wire firebase_messaging for cross-device push notifications, in addition to local notifications.")
    bullet(doc, "Hide the existing 'dashb' key reference in main_screen.dart and verify all bottom-nav labels render correctly.")
    bullet(doc, "Provide a 1024 × 1024 transparent logo with proper safe-zone padding for adaptive icons on Android 8+ launchers.")

    # ---- Appendix: routes & screens ------------------------------------------
    h1(doc, "Appendix A — Route Map")
    table(
        doc,
        ["Path", "Screen", "Notes"],
        [
            ["/", "SplashScreen", "Entry, animation, auth gate."],
            ["/onboarding", "OnboardingScreen", "Shown once until settingsVM.onboardingSeen."],
            ["/login", "LoginScreen", "Combined sign-in & sign-up."],
            ["/home", "MainScreen", "Six-tab bottom nav."],
            ["/add-task", "AddTaskScreen", "Create or edit task (extra = Task?)."],
            ["/settings", "SettingsScreen", "Theme, locale, defaults, restore."],
            ["/help", "HelpScreen", "FAQ + feedback form."],
            ["/pomodoro", "PomodoroScreen", "Focus timer."],
            ["/kanban", "KanbanScreen", "Drag-and-drop board."],
            ["/shared-lists", "SharedListsScreen", "List of collaborative lists."],
            ["/shared-lists/:id/activity", "SharedListActivityScreen", "Audit feed."],
            ["/project/:id", "ProjectDetailScreen", "Tasks within a project."],
            ["/task/:id/reminders", "TaskRemindersScreen", "Multiple reminders editor."],
            ["/notes", "NotesScreen", "Sticky notes grid."],
            ["/search", "SearchScreen", "Cross-collection live search."],
        ],
    )

    h1(doc, "Appendix B — Build & Run Commands")
    code(
        doc,
        "# install dependencies\n"
        "flutter pub get\n\n"
        "# regenerate launcher icons after editing assets/images/logos/logo.png\n"
        "dart run flutter_launcher_icons\n\n"
        "# run on the connected device\n"
        "flutter run\n\n"
        "# release builds\n"
        "flutter build apk --release\n"
        "flutter build ipa --release\n\n"
        "# static analysis\n"
        "flutter analyze\n\n"
        "# export user accounts (admin only)\n"
        "cd scripts/export-users\n"
        "npm install\n"
        "node export_users.js   # writes users.csv\n",
    )

    out = r"D:\ToDoListApp\Focus24_Project_Report.docx"
    try:
        doc.save(out)
    except PermissionError:
        # File is still open in Word — write a timestamped copy instead.
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out = rf"D:\ToDoListApp\Focus24_Project_Report_{stamp}.docx"
        doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
